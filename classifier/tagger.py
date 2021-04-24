import docx
import os
import sys
from google.cloud import language
from google.cloud.language import enums
from google.cloud.language import types
import six

def getText(filename):
    fullText = []
    doc = docx.Document(filename)
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def syntax_text(text):
    """Detects syntax in the text."""
    client = language.LanguageServiceClient()

    if isinstance(text, six.binary_type):
        text = text.decode('utf-8')

    # Instantiates a plain text document.
    document = types.Document(
        content=text,
        type=enums.Document.Type.PLAIN_TEXT)

    # Detects syntax in the document. You can also analyze HTML with:
    # document.type == enums.Document.Type.HTML
    tokens = client.analyze_syntax(document).tokens

    # part-of-speech tags from enums.PartOfSpeech.Tag
    pos_tag = ('UNKNOWN', 'ADJ', 'ADP', 'ADV', 'CONJ', 'DET', 'NOUN', 'NUM',
               'PRON', 'PRT', 'PUNCT', 'VERB', 'X', 'AFFIX')
    texts = []
    pos = []
    for token in tokens:
        texts.append(token.text.content)
        pos.append(pos_tag[token.part_of_speech.tag])
    return texts, pos

def classify(text, verbose=True):
    """Classify the input text into categories. """

    language_client = language.LanguageServiceClient()

    document = language.types.Document(
        content=text,
        type=language.enums.Document.Type.PLAIN_TEXT)
    response = language_client.classify_text(document)
    categories = response.categories

    result = {}

    for category in categories:
        # Turn the categories into a dictionary of the form:
        # {category.name: category.confidence}, so that they can
        # be treated as a sparse vector.
        result[category.name] = category.confidence

    if verbose:
        print(text)
        for category in categories:
            print(u'=' * 20)
            print(u'{:<16}: {}'.format('category', category.name))
            print(u'{:<16}: {}'.format('confidence', category.confidence))

    return result

def entities_text(text):
    """Detects entities in the text."""
    client = language.LanguageServiceClient()

    if isinstance(text, six.binary_type):
        text = text.decode('utf-8')

    # Instantiates a plain text document.
    document = types.Document(
        content=text,
        type=enums.Document.Type.PLAIN_TEXT)

    # Detects entities in the document. You can also analyze HTML with:
    #   document.type == enums.Document.Type.HTML
    entities = client.analyze_entities(document).entities

    # entity types from enums.Entity.Type
    entity_type = ('UNKNOWN', 'PERSON', 'LOCATION', 'ORGANIZATION',
                   'EVENT', 'WORK_OF_ART', 'CONSUMER_GOOD', 'OTHER')

    for entity in entities:
        print('=' * 20)
        print(u'{:<16}: {}'.format('name', entity.name))
        print(u'{:<16}: {}'.format('type', entity_type[entity.type]))
        print(u'{:<16}: {}'.format('metadata', entity.metadata))
        print(u'{:<16}: {}'.format('salience', entity.salience))
        print(u'{:<16}: {}'.format('wikipedia_url',
              entity.metadata.get('wikipedia_url', '-')))

def getCategory(isT):
    if(isT):
        return "T"
    else:
        return "O"


directory = os.fsencode("./Files")
numOfFiles = len(os.listdir(directory))
counter = 0
for file in os.listdir(directory):
    counter += 1
    filename = os.fsdecode(file)
    if filename.endswith(".docx") or filename.endswith(".DOCX"):
        try:
            data_words = []
            data_pos = []
            categories = []
            isT = False
            
            words, pos = syntax_text(getText(filename))
            for i in range(0,len(words)):
                if(words[i] == "START"):
                    isT = True
                if(words[i] == "STOP"):
                    isT = False
                    data_words.append("\n")
                    data_pos.append("newline")
                    categories.append('O')
                if("{" not in words[i] and "}" not in words[i] and "START" not in words[i] and "STOP" not in words[i]):
                    data_words.append(words[i])
                    data_pos.append(pos[i])
                    categories.append(getCategory(isT))
            
            finalText = ""
            newLine = False
            for i in range(0,len(data_words)):
                word = data_words[i]
                category = categories[i]
                if(category == "T"):
                    finalText += word.replace("\n", "") + " "
                    newLine = True
                else:
                    if(newLine):
                        finalText += "\n\n"
                        newLine = False
          
            if(finalText != "\n" and finalText != "\n\n"):
                file = open("structureData.csv","a")
                for i in range(0,len(data_words)):
                    word = data_words[i]
                    category = categories[i]
                    ps = pos[i]
                    word = word.replace("\"","")
                    if(word.replace("\n", "") != "" and word.replace("\n", "") != " "):
                        file.write("\"" + word.replace("\n", "") + "\"," + ps + "," + category + "\n")
            print("Closing" + str(filename) + " | File #" + str(counter) + " of " + str(numOfFiles))
            file.close()
        except:
            print("Error with: " + filename)
            print()
            print(os.path.abspath('files/' + filename))
            print()
            print(sys.exc_info())
            print("----------------------")
