import sys
import docx
#sys.stdout = open('output.txt', 'w')
import nltk
#nltk.download('punkt')
from nltk.stem.lancaster import LancasterStemmer
import numpy as np
import tflearn
import tensorflow as tf
import csv
import unicodedata
from google.cloud import language
from google.cloud.language import enums
from google.cloud.language import types
import six

def syntax_text(text):
    """Detects syntax in the text."""
    client = language.LanguageServiceClient()

    if isinstance(text, six.binary_type):
        text = text.decode('utf-8')

    # Instantiates a plain text document.
    document = types.Document(
        content=text,
        type=enums.Document.Type.PLAIN_TEXT)

    # Detects syntax in the document. You can also analyze HTML with:
    # document.type == enums.Document.Type.HTML
    tokens = client.analyze_syntax(document).tokens

    # part-of-speech tags from enums.PartOfSpeech.Tag
    pos_tag = ('UNKNOWN', 'ADJ', 'ADP', 'ADV', 'CONJ', 'DET', 'NOUN', 'NUM',
               'PRON', 'PRT', 'PUNCT', 'VERB', 'X', 'AFFIX')
    texts = []
    pos = []
    for token in tokens:
        texts.append(token.text.content)
        pos.append(pos_tag[token.part_of_speech.tag])
    return texts, pos

# a table structure to hold the different punctuation used
tbl = dict.fromkeys(i for i in range(sys.maxunicode)
                    if unicodedata.category(chr(i)).startswith('P'))

# method to remove punctuations from sentences.
def remove_punctuation(text):
    return text.translate(tbl)


# initialize the stemmer
stemmer = LancasterStemmer()

words = []
categories = ['O','T']
docs = []
with open('structureData.csv', 'rt') as csvfile:
     inputs = csv.reader(csvfile, delimiter=',', quotechar='"')
     for inp in inputs:
         docs.append([inp[0],inp[2]])
         words.append(inp[0])

# stem and lower each word and remove duplicates
words = [stemmer.stem(w.lower()) for w in words]
words = sorted(list(set(words)))


# create our training data
training = []
output = []
# create an empty array for our output
output_empty = [0] * len(categories)

for doc in docs:
    # initialize our bag of words(bow) for each document in the list
    bow = []
    # list of tokenized words for the pattern
    token_words = doc[0]
    # stem each word
    token_words = [stemmer.stem(word.lower()) for word in token_words]

    # create our bag of words array
    for w in words:
        bow.append(1) if w in token_words else bow.append(0)
    
    output_row = list(output_empty)
    output_row[categories.index(doc[1])] = 1

    # our training set will contain a the bag of words model and the output row that tells
    # which catefory that bow belongs to.
    training.append([bow, output_row])

# shuffle our features and turn into np.array as tensorflow  takes in numpy array
training = np.array(training)
print(training)

# trainX contains the Bag of words and train_y contains the label/ category
train_x = list(training[:, 0])
train_y = list(training[:, 2])

# reset underlying graph data
tf.reset_default_graph()
# Build neural network
net = tflearn.input_data(shape=[None, len(train_x[0])])
net = tflearn.fully_connected(net, 8)
net = tflearn.fully_connected(net, 8)
net = tflearn.fully_connected(net, len(train_y[0]), activation='softmax')
net = tflearn.regression(net)

# Define model and setup tensorboard
model = tflearn.DNN(net, tensorboard_dir='tflearn_logs', tensorboard_verbose=3)
# Start training (apply gradient descent algorithm)
model.fit(train_x, train_y, n_epoch=100, batch_size=8, show_metric=True)
model.save('model.tflearn')

def get_tf_record(sentence):
    global words
    # tokenize the pattern
    sentence_words = nltk.word_tokenize(sentence)
    # stem each word
    sentence_words = [stemmer.stem(word.lower()) for word in sentence_words]
    # bag of words
    bow = [0]*len(words)
    for s in sentence_words:
        for i, w in enumerate(words):
            if w == s:
                bow[i] = 1

    return(np.array(bow))

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
"""
for word in syntax_text(getText('file.docx'))[0]:
    print(word + "||" + categories[np.argmax(model.predict([get_tf_record(doc[0])]))])
"""





